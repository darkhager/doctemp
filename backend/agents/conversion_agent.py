"""Team Alpha — Core Engine: DOCX ↔ template bi-directional conversion."""
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from schemas import FieldSchema

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


class ConversionAgent:
    def docx_to_html(self, docx_path: str | Path) -> tuple[str, list[FieldSchema]]:
        """
        Parse a .docx file and return (html_content, detected_fields).
        Placeholders in the form {{name}} are preserved; other text is wrapped in <p> tags.
        """
        doc = Document(str(docx_path))
        html_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                html_parts.append("<p><br></p>")
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                html_parts.append(f"<h1>{self._escape(text)}</h1>")
            elif "Heading 2" in style:
                html_parts.append(f"<h2>{self._escape(text)}</h2>")
            elif "Heading 3" in style:
                html_parts.append(f"<h3>{self._escape(text)}</h3>")
            else:
                inline = self._runs_to_html(para)
                html_parts.append(f"<p>{inline}</p>")

        for table in doc.tables:
            rows_html = []
            for i, row in enumerate(table.rows):
                cells = "".join(
                    f"<{'th' if i == 0 else 'td'}>{self._escape(cell.text)}<{'th' if i == 0 else '/td'}>"
                    for cell in row.cells
                )
                rows_html.append(f"<tr>{cells}</tr>")
            html_parts.append(f"<table>{''.join(rows_html)}</table>")

        html = "\n".join(html_parts)
        names = list(dict.fromkeys(PLACEHOLDER_RE.findall(html)))
        fields = [
            FieldSchema(
                name=n,
                label=n.replace("_", " ").title(),
                field_type="text",
                required=False,
                default_value="",
            )
            for n in names
        ]
        return html, fields

    def fill_template(self, html: str, values: dict[str, str]) -> str:
        """Replace {{key}} placeholders with provided values."""
        def replacer(m: re.Match) -> str:
            return values.get(m.group(1), m.group(0))
        return PLACEHOLDER_RE.sub(replacer, html)

    def _runs_to_html(self, para) -> str:
        parts = []
        for run in para.runs:
            text = self._escape(run.text)
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            parts.append(text)
        return "".join(parts) or self._escape(para.text)

    @staticmethod
    def _escape(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            # restore placeholders that got escaped
            .replace("&lt;&lt;", "{{")
            .replace("&gt;&gt;", "}}")
        )
