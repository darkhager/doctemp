"""Team Alpha — Core Engine: export filled templates to .docx or .pdf.

The editor stores documents as a sequence of pages:

    <div class="doc-page" data-orientation="portrait|landscape">
        <div class="doc-sec" data-type="header|body|footer" data-label="...">…</div>
        …
    </div>

For .docx each page becomes its own Word *section* sized to A4 with the page's
orientation, so the result opens in Microsoft Word with correct page geometry.
Header/footer sections map to real Word header/footer regions; body sections flow
into the page body.
"""
from io import BytesIO
from html.parser import HTMLParser
from urllib.parse import unquote

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION

VOID = {"br", "img", "hr", "input", "meta", "link"}


def _fmt_start(tag, attrs, selfclose=False):
    a = "".join(f' {k}="{v}"' if v is not None else f" {k}" for k, v in attrs)
    return f"<{tag}{a}{'/' if selfclose else ''}>"


class _PageSegmenter(HTMLParser):
    """Split the document HTML into pages → sections, capturing each section's inner HTML.

    Handles arbitrarily nested <div>s inside a section (e.g. panel grids) by tracking
    div depth rather than relying on regex.
    """
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.pages: list[dict] = []
        self._depth = 0
        self._page = None
        self._page_depth = None
        self._sec = None
        self._sec_depth = None
        self._buf: list[str] = []

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        cls = a.get("class", "") or ""
        if tag == "div" and "doc-page" in cls:
            self._page = {"orientation": a.get("data-orientation") or "portrait", "sections": []}
            self.pages.append(self._page)
            self._page_depth = self._depth
            self._depth += 1
            return
        if tag == "div" and "doc-sec" in cls and self._page is not None:
            self._sec = {"type": a.get("data-type") or "body",
                         "label": unquote(a.get("data-label") or ""),
                         "html": ""}
            self._page["sections"].append(self._sec)
            self._sec_depth = self._depth
            self._buf = []
            self._depth += 1
            return
        if self._sec is not None:
            self._buf.append(_fmt_start(tag, attrs))
        if tag not in VOID:
            self._depth += 1

    def handle_startendtag(self, tag, attrs):
        if self._sec is not None:
            self._buf.append(_fmt_start(tag, attrs, selfclose=True))

    def handle_endtag(self, tag):
        if tag in VOID:
            return
        self._depth -= 1
        if self._sec is not None and self._depth == self._sec_depth:
            self._sec["html"] = "".join(self._buf)
            self._sec = None
            self._sec_depth = None
            self._buf = []
            return
        if self._page is not None and self._sec is None and self._depth == self._page_depth:
            self._page = None
            self._page_depth = None
            return
        if self._sec is not None:
            self._buf.append(f"</{tag}>")

    def handle_data(self, data):
        if self._sec is not None:
            self._buf.append(data)


class _BlockParser(HTMLParser):
    """Convert a section's inner HTML into a flat list of blocks for docx output.

    Blocks: {'kind': 'para'|'heading'|'bullet', 'level', 'runs': [(text, b, i, u)]}
            {'kind': 'table', 'rows': [[cell, …], …]}
    """
    HEADINGS = {"h1": 1, "h2": 2, "h3": 3}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks: list[dict] = []
        self._runs: list[tuple] = []
        self._kind = None
        self._level = 0
        self._b = self._i = self._u = False
        self._intable = False
        self._table = None
        self._row = None
        self._cell = None

    def _flush(self):
        if self._kind and self._runs:
            self.blocks.append({"kind": self._kind, "level": self._level, "runs": self._runs})
        self._runs, self._kind, self._level = [], None, 0

    def handle_starttag(self, tag, attrs):
        if tag in ("strong", "b"): self._b = True; return
        if tag in ("em", "i"):     self._i = True; return
        if tag == "u":             self._u = True; return
        if tag == "br":            self._runs.append(("\n", self._b, self._i, self._u)); return
        if tag == "img":
            if self._kind is None: self._kind = "para"
            self._runs.append(("[Image]", self._b, self._i, self._u)); return
        if tag in self.HEADINGS:
            self._flush(); self._kind = "heading"; self._level = self.HEADINGS[tag]; return
        if tag == "li":
            self._flush(); self._kind = "bullet"; return
        if tag in ("p", "div"):
            self._flush(); self._kind = "para"; return
        if tag == "table":
            self._flush(); self._intable = True; self._table = []; return
        if tag == "tr" and self._intable: self._row = []; return
        if tag in ("td", "th") and self._intable: self._cell = []; return

    def handle_endtag(self, tag):
        if tag in ("strong", "b"): self._b = False; return
        if tag in ("em", "i"):     self._i = False; return
        if tag == "u":             self._u = False; return
        if tag in self.HEADINGS or tag in ("li", "p", "div"):
            self._flush(); return
        if tag in ("td", "th") and self._intable:
            self._row.append("".join(self._cell or []).strip()); self._cell = None; return
        if tag == "tr" and self._intable and self._row is not None:
            self._table.append(self._row); self._row = None; return
        if tag == "table" and self._intable:
            self.blocks.append({"kind": "table", "rows": self._table})
            self._intable = False; self._table = None; return

    def handle_data(self, data):
        if self._intable:
            if self._cell is not None: self._cell.append(data)
            return
        if not data.strip() and self._kind is None:
            return
        if self._kind is None: self._kind = "para"
        self._runs.append((data, self._b, self._i, self._u))


def _segment_pages(html: str) -> list[dict]:
    p = _PageSegmenter()
    p.feed(html or "")
    return p.pages


def _blocks(html: str) -> list[dict]:
    p = _BlockParser()
    p.feed(html or "")
    p._flush()
    return p.blocks


_HEADING_PT = {1: 18, 2: 15, 3: 13}


class RenderAgent:
    def to_docx(self, content_html: str, field_values: dict[str, str] = None) -> bytes:
        """Generate a Word .docx: one A4 section per page, per-page orientation,
        header/footer mapped to Word header/footer regions."""
        from agents.conversion_agent import ConversionAgent
        filled = ConversionAgent().fill_template(content_html, field_values or {})

        pages = _segment_pages(filled)
        doc = Document()

        if not pages:
            # No page structure (legacy/plain HTML) — single portrait A4 page.
            self._apply_geometry(doc.sections[0], "portrait")
            self._write_blocks(doc, _blocks(filled), allow_table=True)
            return self._save(doc)

        for idx, page in enumerate(pages):
            section = doc.sections[0] if idx == 0 else doc.add_section(WD_SECTION.NEW_PAGE)
            self._apply_geometry(section, page["orientation"])
            for sec in page["sections"]:
                blocks = _blocks(sec["html"])
                if not blocks:
                    continue
                if sec["type"] == "header":
                    self._unlink(section.header)
                    self._write_blocks(section.header, blocks, allow_table=False)
                elif sec["type"] == "footer":
                    self._unlink(section.footer)
                    self._write_blocks(section.footer, blocks, allow_table=False)
                else:
                    self._write_blocks(doc, blocks, allow_table=True)

        return self._save(doc)

    def to_pdf(self, content_html: str, field_values: dict[str, str] = None) -> bytes:
        """Generate a filled PDF. A4 pages; each editor page on its own PDF page."""
        from agents.conversion_agent import ConversionAgent
        filled = ConversionAgent().fill_template(content_html, field_values or {})
        full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm; }}
  body {{ font-family: Arial, sans-serif; font-size: 12pt; }}
  h1 {{ font-size: 20pt; }} h2 {{ font-size: 16pt; }} h3 {{ font-size: 13pt; }}
  table {{ border-collapse: collapse; width: 100%; }}
  td, th {{ border: 1px solid #ccc; padding: 6px 10px; }}
  th {{ background: #f0f0f0; font-weight: bold; }}
  .doc-page {{ page-break-after: always; }}
  .doc-page:last-child {{ page-break-after: auto; }}
</style>
</head>
<body>{filled}</body>
</html>"""
        from weasyprint import HTML
        return HTML(string=full_html).write_pdf()

    # ── helpers ────────────────────────────────────────────────────────────────

    @staticmethod
    def _apply_geometry(section, orientation: str) -> None:
        """Size a Word section to A4 with the given orientation (Word needs the
        width/height swapped manually — setting orientation alone does not)."""
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = Mm(297), Mm(210)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = Mm(210), Mm(297)
        section.left_margin = section.right_margin = Mm(25.4)
        section.top_margin = section.bottom_margin = Mm(25.4)

    @staticmethod
    def _unlink(part) -> None:
        try:
            part.is_linked_to_previous = False
        except Exception:
            pass

    @staticmethod
    def _write_blocks(container, blocks: list[dict], allow_table: bool) -> None:
        for blk in blocks:
            if blk["kind"] == "table":
                rows = blk.get("rows") or []
                cols = max((len(r) for r in rows), default=0)
                if not cols:
                    continue
                if allow_table:
                    table = container.add_table(rows=0, cols=cols)
                    try:
                        table.style = "Table Grid"
                    except Exception:
                        pass
                    for r in rows:
                        cells = table.add_row().cells
                        for ci, val in enumerate(r):
                            if ci < cols:
                                cells[ci].text = val
                else:
                    # header/footer can't host tables cleanly — flatten to text rows.
                    for r in rows:
                        container.add_paragraph(" | ".join(r))
                continue

            p = container.add_paragraph()
            if blk["kind"] == "bullet":
                try:
                    p.style = "List Bullet"
                except Exception:
                    pass
            is_heading = blk["kind"] == "heading"
            for text, b, i, u in blk.get("runs", []):
                if text == "\n":
                    p.add_run().add_break()
                    continue
                run = p.add_run(text)
                run.bold = b or is_heading
                run.italic = i
                run.underline = u
                if is_heading:
                    run.font.size = Pt(_HEADING_PT.get(blk["level"], 14))

    @staticmethod
    def _save(doc) -> bytes:
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
